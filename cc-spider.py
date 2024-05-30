import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from datetime import datetime
import os,time

#初始化变量
DEBUG = True
USER = '葡萄'                               #配置用户名
SLEEP_TIME = 5                              #抓取每篇文章的时间间隔，单位秒
start_page = 874                            #起始页面
end_page   = 1007                           #结束页面
FIXED_URL  = 'https://talkcc.org/'          #固定URL
FILE_NAME  = f'./西西河-{USER}-专辑.docx'    #生成word文档名 
LOG_FILE   = f'./cc-spider.log'             #生成日志文件名
         

# 初始化Word文档
if os.path.exists(FILE_NAME):
    # 文件存在，打开并追加内容
    doc = Document(FILE_NAME)
    doc.add_paragraph('')
    doc.save(FILE_NAME)
else:
    # 文件不存在，创建新文件
    doc = Document()
    doc.add_heading(f"{USER}文章集合", 0)
    doc.save(FILE_NAME)

def log_message(message):
    current_time = datetime.now()
    formatted_time = current_time.strftime("%Y-%m-%d %H:%M:%S  ")
    if DEBUG:
        print(message)  #output to terminal
    
    with open(LOG_FILE, 'a') as file:
        # 将内容追加到文件中
        file.write(formatted_time + message + '\n')



def http_req(base_url):
    headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'
    }

    # 设置cookies，注意URL解码
    cc_value = 'DcuxFcMgDEXRhVRISF%2BCaTgYmU1SZA%2BXLtNmm3iP0LzqPoSo9uf1fq7P73v3YswnMkaKzTFCQ5srN6CMlrN1CXGDSw0ycaftd%2FdmhFIbFSGmwsR9yEpkRRxTXAXuinPVwzwRNlfnPw%3D%3D'
    cookies = {
        'cc': requests.utils.unquote(cc_value)
    }
    # 发送请求并解析网页以获取文章链接
    response = requests.get(base_url, headers=headers, cookies=cookies) 
    return response

def get_article_links(page_rsp):
    soup = BeautifulSoup(page_rsp.text, 'html.parser')
    div_tree = soup.find('div', class_='s_Tree')    #format <div class="s_Tree">

    all_links = div_tree.find_all('a', href=True)  # format: <a href="/article/4487276">xxxx</a>
    all_dates = div_tree.find_all('small')  #format: <small>2020-03-08 12:43:34</small>

    articles_links = [s for s in all_links if 'article' in str(s)]
    return all_dates[::-1], articles_links[::-1]     #reverse the list for correct order

def handle_article(url, date, article_response):
    article_soup = BeautifulSoup(article_response.text, 'html.parser')
    #print(article_soup)

    # 提取文章内容，这里需要根据实际页面结构来调整
    sections = article_soup.find(lambda tag: tag.name == 'div' and tag.get('class') == ['s_Sec']) # 精确匹配 <div class="s_Sec">
    if sections is None:
        log_message(f'ERROR: when reading article {url}')
        return
    title = sections.find_all('b')
    p_tags = sections.find_all('p')

    # 将文章标题和内容添加到Word文档
    valid_idx = 1;
    doc.add_heading(title[valid_idx].get_text(), level=3)
    doc.add_paragraph(f'原文：{url}')
    doc.add_paragraph(date)
    for p_tag in p_tags:
        if not p_tag.findChildren():    #no subnodes and only string
            doc.add_paragraph(p_tag)     
        elif p_tag.find('a'):     #link
            h_ref = p_tag.a['href']
            doc.add_paragraph(f'{p_tag.a.string} ({h_ref})')
        elif p_tag.find('img'):   #pictures
            img_element = p_tag.find('img')                 
            src_value = img_element['src']      #<img ... src="/picture/0,2405/2256_06203932.png"/>
            img_path = download_image(FIXED_URL+src_value, './'+src_value)

            if img_path == None:
                doc.add_paragraph('missing img src from ' + FIXED_URL + src_value)
            else:
                doc.add_picture(img_path, width=Inches(5.5))
        else:
            doc.add_paragraph(p_tag.get_text())
            log_message(f"INFO: handle special p_tag = {p_tag}")
                       
    # 可以选择添加一些分隔符或格式化
    doc.add_paragraph('')
    doc.add_paragraph('')
    #doc.add_page_break()

def download_image(url, file_path):
    # 兼容Windows/linux风格的路径
    file_path = os.path.join(*file_path.split('/'))
    # 创建目录（如果不存在）
    os.makedirs(os.path.dirname(file_path), exist_ok=True)

    response = requests.get(url)
    if response.status_code == 200:
        with open(file_path, 'wb') as file:
            file.write(response.content)
        log_message(f"    Image downloaded and saved as {file_path}")
        return file_path
    else:
        log_message(f"ERROR: Failed to download image from {url}")
        return None

def print_article(date, url):
    article_response = http_req(url)
    if (article_response.status_code == 200):
        handle_article(url, date, article_response)
    else:
        log_message(f'ERROR: fail to get article from {url}')

# 目标网址
def main():
    HOME_PATH = f'{FIXED_URL}user/{USER}/所有帖/'
    log_message('-----------------------------------------------------------------------------------')
    log_message(f'Ready to collect {USER}\'s articles from {HOME_PATH}, page {start_page}~{end_page}')
    
    start_time = time.time()    #record the start time
    total_articles = 0          #Initialize a couter for total articles

    for page in range(start_page, end_page+1): 
        log_message(f'Start to process Page-{page}')
        
        page_url = f'{HOME_PATH}{page}'
        page_rsp = http_req(page_url)
        if (page_rsp.status_code != 200):
            log_message(f'ERROR: fail to get Page-{page}，status_code：{page_rsp.status_code}')
            continue

        #page rsp successful
        dates, articles_links = get_article_links(page_rsp)
        
        # 循环处理每个文章链接
        for i, link in enumerate(articles_links):
            article_url = link['href']
            if article_url.startswith('/article/'):
                log_message(f'    Start to process {article_url}')
                url = f'{FIXED_URL}{article_url}'
                print_article(dates[i], url)
                total_articles += 1
                time.sleep(SLEEP_TIME)
        
        doc.save(FILE_NAME)  #save for every page

    #Finally save the file
    doc.save(FILE_NAME)
    
    end_time = time.time()
    elapsed_time = end_time - start_time
    hours =  int(elapsed_time // 3600)
    minutes = int((elapsed_time % 3600) // 60)
    seconds = int(elapsed_time % 60)

    log_message(f'Done! Total articles downloaded: {total_articles}')
    log_message(f'Elapsed time: {hours:02d}:{minutes:02d}:{seconds:02d}s')
    
    
if __name__ == "__main__":
    if 1:
        main()
    else:   #test code
        print_article('', 'https://talkcc.org/article/4108278')
